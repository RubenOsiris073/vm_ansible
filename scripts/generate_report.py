#!/usr/bin/env python3
"""
Script para generar reporte en formato DOCX basado en el repositorio vm_ansible
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import subprocess

def add_heading(doc, text, level=1):
    """Agregar encabezado con formato personalizado"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph_with_format(doc, text, bold=False, italic=False):
    """Agregar párrafo con formato"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_code_block(doc, code_text):
    """Agregar bloque de código con formato monoespaciado"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Fondo gris claro simulado con tabulación
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def parse_markdown_content(content):
    """Parsear contenido markdown y retornar secciones estructuradas"""
    sections = []
    current_section = None
    current_content = []
    in_code_block = False
    code_block_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        # Detectar inicio/fin de bloque de código
        if line.strip().startswith('```'):
            if in_code_block:
                # Fin del bloque de código
                sections.append({
                    'type': 'code',
                    'content': '\n'.join(code_block_content)
                })
                code_block_content = []
                in_code_block = False
            else:
                # Inicio del bloque de código
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Detectar encabezados
        if line.startswith('#'):
            # Guardar sección anterior si existe
            if current_section and current_content:
                sections.append({
                    'type': 'section',
                    'title': current_section['title'],
                    'level': current_section['level'],
                    'content': '\n'.join(current_content)
                })
                current_content = []
            
            # Nueva sección
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('#').strip()
            current_section = {'title': title, 'level': level}
        elif current_section:
            # Contenido de la sección actual
            if line.strip():
                current_content.append(line)
    
    # Agregar última sección
    if current_section and current_content:
        sections.append({
            'type': 'section',
            'title': current_section['title'],
            'level': current_section['level'],
            'content': '\n'.join(current_content)
        })
    
    return sections

def get_git_info():
    """Obtener información del repositorio git"""
    try:
        branch = subprocess.check_output(['git', 'rev-parse', '--abbrev-ref', 'HEAD'], 
                                        text=True).strip()
        commit = subprocess.check_output(['git', 'rev-parse', '--short', 'HEAD'], 
                                        text=True).strip()
        return f"Rama: {branch}, Commit: {commit}"
    except:
        return "Información de git no disponible"

def get_directory_structure(path, prefix='', max_depth=3, current_depth=0):
    """Obtener estructura de directorios recursivamente"""
    if current_depth >= max_depth:
        return []
    
    structure = []
    try:
        items = sorted(os.listdir(path))
        # Filtrar archivos ocultos y directorios no deseados
        items = [item for item in items if not item.startswith('.') 
                 and item not in ['node_modules', '__pycache__', 'venv']]
        
        for i, item in enumerate(items):
            item_path = os.path.join(path, item)
            is_last = i == len(items) - 1
            connector = '└── ' if is_last else '├── '
            
            if os.path.isdir(item_path):
                structure.append(f"{prefix}{connector}{item}/")
                extension = '    ' if is_last else '│   '
                structure.extend(get_directory_structure(
                    item_path, 
                    prefix + extension, 
                    max_depth, 
                    current_depth + 1
                ))
            else:
                structure.append(f"{prefix}{connector}{item}")
    except PermissionError:
        pass
    
    return structure

def generate_report(repo_path, output_file):
    """Generar reporte completo en DOCX"""
    print(f"Generando reporte desde: {repo_path}")
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Portada
    add_heading(doc, 'RouterLab - Sistema de Gestión de Red', 1)
    add_heading(doc, 'Reporte Técnico del Proyecto', 2)
    
    # Información del documento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\nFecha de Generación: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}\n')
    run.italic = True
    
    git_info = get_git_info()
    run = p.add_run(f'{git_info}\n')
    run.italic = True
    run.font.size = Pt(10)
    
    # Salto de página
    doc.add_page_break()
    
    # Índice manual (Word puede generar uno automático)
    add_heading(doc, 'Contenido', 1)
    doc.add_paragraph('1. Introducción')
    doc.add_paragraph('2. Arquitectura del Sistema')
    doc.add_paragraph('3. Estructura del Proyecto')
    doc.add_paragraph('4. Componentes Principales')
    doc.add_paragraph('5. Guía de Inicio Rápido')
    doc.add_paragraph('6. Configuración y Despliegue')
    doc.add_paragraph('7. APIs del Sistema')
    doc.add_paragraph('8. Referencias')
    
    doc.add_page_break()
    
    # Leer y procesar README.md
    readme_path = os.path.join(repo_path, 'README.md')
    if os.path.exists(readme_path):
        print("Procesando README.md...")
        with open(readme_path, 'r', encoding='utf-8') as f:
            readme_content = f.read()
        
        sections = parse_markdown_content(readme_content)
        
        for section in sections:
            if section['type'] == 'section':
                add_heading(doc, section['title'], min(section['level'], 3))
                # Limpiar y agregar contenido
                content = section['content'].strip()
                if content:
                    # Dividir en párrafos
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            # Detectar listas
                            if para.strip().startswith('-') or para.strip().startswith('*'):
                                lines = para.split('\n')
                                for line in lines:
                                    if line.strip():
                                        doc.add_paragraph(line.strip().lstrip('-*').strip(), 
                                                        style='List Bullet')
                            else:
                                doc.add_paragraph(para.strip())
            
            elif section['type'] == 'code':
                add_code_block(doc, section['content'])
    
    # Agregar estructura de directorios
    doc.add_page_break()
    add_heading(doc, 'Estructura Completa del Repositorio', 1)
    doc.add_paragraph('Estructura de directorios y archivos del proyecto:')
    doc.add_paragraph()
    
    structure = get_directory_structure(repo_path)
    if structure:
        structure_text = '\n'.join(structure)
        add_code_block(doc, structure_text)
    
    # Leer y procesar SSH_k3s_quick.md si existe
    ssh_guide_path = os.path.join(repo_path, 'SSH_k3s_quick.md')
    if os.path.exists(ssh_guide_path):
        print("Procesando SSH_k3s_quick.md...")
        doc.add_page_break()
        add_heading(doc, 'Guía Rápida SSH y K3s', 1)
        
        with open(ssh_guide_path, 'r', encoding='utf-8') as f:
            ssh_content = f.read()
        
        sections = parse_markdown_content(ssh_content)
        for section in sections:
            if section['type'] == 'section':
                add_heading(doc, section['title'], min(section['level'] + 1, 3))
                if section['content'].strip():
                    doc.add_paragraph(section['content'].strip())
            elif section['type'] == 'code':
                add_code_block(doc, section['content'])
    
    # Información de componentes clave
    doc.add_page_break()
    add_heading(doc, 'Inventario de Componentes', 1)
    
    # Apps
    add_heading(doc, 'Aplicaciones', 2)
    apps_path = os.path.join(repo_path, 'apps')
    if os.path.exists(apps_path):
        for app in os.listdir(apps_path):
            app_path = os.path.join(apps_path, app)
            if os.path.isdir(app_path):
                doc.add_paragraph(f'• {app}', style='List Bullet')
                # Leer package.json si existe
                package_json = os.path.join(app_path, 'package.json')
                if os.path.exists(package_json):
                    with open(package_json, 'r') as f:
                        import json
                        try:
                            pkg = json.load(f)
                            if 'description' in pkg:
                                p = doc.add_paragraph()
                                p.add_run(f"  Descripción: {pkg['description']}").italic = True
                        except:
                            pass
    
    # Playbooks
    add_heading(doc, 'Playbooks de Ansible', 2)
    playbooks_path = os.path.join(repo_path, 'playbooks')
    if os.path.exists(playbooks_path):
        for playbook in os.listdir(playbooks_path):
            if playbook.endswith('.yml') or playbook.endswith('.yaml'):
                doc.add_paragraph(f'• {playbook}', style='List Bullet')
    
    # Roles
    add_heading(doc, 'Roles de Ansible', 2)
    roles_path = os.path.join(repo_path, 'roles')
    if os.path.exists(roles_path):
        for role in os.listdir(roles_path):
            role_path = os.path.join(roles_path, role)
            if os.path.isdir(role_path):
                doc.add_paragraph(f'• {role}', style='List Bullet')
    
    # Conclusión
    doc.add_page_break()
    add_heading(doc, 'Conclusión', 1)
    doc.add_paragraph(
        'Este documento proporciona una visión completa del proyecto RouterLab, '
        'incluyendo su arquitectura, componentes, configuración y guías de uso. '
        'El sistema está diseñado para gestionar y monitorear dispositivos de red '
        'utilizando tecnologías modernas como Kubernetes, Ansible y Docker.'
    )
    
    # Guardar documento
    doc.save(output_file)
    print(f"✓ Reporte generado exitosamente: {output_file}")

def main():
    # Determinar ruta del repositorio
    script_dir = os.path.dirname(os.path.abspath(__file__))
    repo_path = os.path.dirname(script_dir)
    
    # Nombre del archivo de salida
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(repo_path, f'RouterLab_Reporte_{timestamp}.docx')
    
    print("=" * 60)
    print("GENERADOR DE REPORTE DOCX - RouterLab")
    print("=" * 60)
    
    generate_report(repo_path, output_file)
    
    print("=" * 60)
    print(f"Reporte guardado en: {output_file}")
    print("=" * 60)

if __name__ == '__main__':
    main()
